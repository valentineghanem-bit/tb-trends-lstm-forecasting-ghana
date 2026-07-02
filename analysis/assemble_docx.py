"""
assemble_docx.py — Stage 10: assemble Q1 manuscript. Vancouver re-index by first appearance,
embed figures AND formatted tables, full front/back matter. Manuscript .docx is a deliverable
(NOT committed to git, Tenet 20).
"""
import re, os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import pandas as pd
from _labels import relabel_cols
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MAN = "manuscript/"
FIG = "outputs/figures/"
TAB = "outputs/tables/"

REFS = {
    "Boah et al. 2021": "Boah M, Kpordoxah MR, Adokiya MN. Self-reported gender differentials in the knowledge of tuberculosis transmission and curative possibility using national representative data in Ghana. PLoS One. 2021;16(7):e0254499. doi:10.1371/journal.pone.0254499.",
    "Osei et al. 2023": "Osei E, Amu H, Kye-Duodu G, Kwabla MP, Danso E, Binka FN, Kim SY. Impact of COVID-19 pandemic on Tuberculosis and HIV services in Ghana: An interrupted time series analysis. PLoS One. 2023;18(9):e0291808. doi:10.1371/journal.pone.0291808.",
    "McQuaid et al. 2022": "McQuaid CF, Henrion MYR, Burke RM, MacPherson P, Nzawa-Soko R, Horton KC. Inequalities in the impact of COVID-19-associated disruptions on tuberculosis diagnosis by age and sex in 45 high TB burden countries. BMC Med. 2022;20(1):432. doi:10.1186/s12916-022-02624-6.",
    "Falzon et al. 2023": "Falzon D, et al. The impact of the COVID-19 pandemic on the global tuberculosis epidemic. Front Immunol. 2023.",
    "World Health Organization 2024": "World Health Organization. Global Tuberculosis Report 2024. Geneva: World Health Organization; 2024.",
    "Lungu et al. 2022": "Lungu P, et al. Effects of COVID-19 pandemic and response on tuberculosis case-finding and treatment, Zambia. Bull World Health Organ. 2022. PMID:35261409.",
    "Soko et al. 2021": "Soko RN, et al. Effects of Coronavirus Disease Pandemic on Tuberculosis Notifications, Malawi. Emerg Infect Dis. 2021.",
    "Pontes et al. 2024": "Pontes CFS, et al. Quantifying disruptions to tuberculosis case detection during the COVID-19 pandemic, Brazil. Rev Inst Med Trop Sao Paulo. 2024.",
    "Dlangalala et al. 2023": "Dlangalala T, et al. Impact of COVID-19 on tuberculosis case investigations and confirmed cases, eThekwini, South Africa. Sci Rep. 2023.",
    "Panford et al. 2022": "Panford V, Kumah E, Kokuro C, et al. Treatment outcomes and associated factors among patients with multidrug-resistant tuberculosis in Ashanti Region, Ghana: a retrospective, cross-sectional study. BMJ Open. 2022;12(7):e062857. doi:10.1136/bmjopen-2022-062857.",
    "Abbew et al. 2025": "Abbew ET, Laryea R, Kwakye AO, et al. Treatment outcomes of multi-drug-resistant and rifampicin-resistant tuberculosis with and without isolation of nontuberculous mycobacteria between 2018-2021: A retrospective cohort study in Ghana. PLoS Negl Trop Dis. 2025;19(7):e0013204. doi:10.1371/journal.pntd.0013204.",
    "Otoo et al. 2025": "Otoo DM, Efichie E, Aborbor DL, et al. Understanding predictors of medication adherence and treatment outcomes among TB patients in the Western Region, Ghana. J Health Popul Nutr. 2025;44(1):407. doi:10.1186/s41043-025-01145-1.",
    "Trébucq et al. 2017": "Trébucq A, Schwoebel V, Kashongwe Z, et al. Treatment outcome with a short multidrug-resistant tuberculosis regimen in nine African countries. Int J Tuberc Lung Dis. 2017;22(1):17-25. doi:10.5588/ijtld.17.0498.",
    "Toft et al. 2022": "Toft AL, Dahl VN, Sifna A, et al. Treatment outcomes for multidrug- and rifampicin-resistant tuberculosis in Central and West Africa: a systematic review and meta-analysis. Int J Infect Dis. 2022;124(Suppl 1):S107-S116. doi:10.1016/j.ijid.2022.08.015.",
    "Mohammed et al. 2024": "Mohammed A, Cheabu BSN, Amoah-Larbi J, et al. Perception and experience of HIV-induced stigma among people with HIV seeking healthcare in Ghana. BMC Health Serv Res. 2024;24(1):1449. doi:10.1186/s12913-024-11930-z.",
    "Yaidoo et al. 2026": "Yaidoo S, Yelbert JMT, MacScott-Lutterodt R, et al. Integrated management of multidrug-resistant tuberculosis, HIV, and hepatitis B co-infection in Ghana: a case report. Front Med (Lausanne). 2026;13:1801109. doi:10.3389/fmed.2026.1801109.",
    "Abade et al. 2024": "Abade A, Porto LF, Scholze AR, et al. A comparative analysis of classical and machine learning methods for forecasting TB/HIV co-infection. Sci Rep. 2024;14(1):18991. doi:10.1038/s41598-024-69580-4.",
    "Yang et al. 2022": "Yang E, Zhang H, Guo X, Zang Z, Liu Z, Liu Y. A multivariate multi-step LSTM forecasting model for tuberculosis incidence with model explanation in Liaoning Province, China. BMC Infect Dis. 2022;22(1):490. doi:10.1186/s12879-022-07462-8.",
    "Wang et al. 2018": "Wang H, Tian CW, Wang WM, Luo XM. Time-series analysis of tuberculosis from 2005 to 2017 in China. Epidemiol Infect. 2018;146(8):935-939. doi:10.1017/S0950268818001115.",
    "von Elm et al. 2007": "von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. Ann Intern Med. 2007;147(8):573-7.",
    "Benchimol et al. 2015": "Benchimol EI, Smeeth L, Guttmann A, et al. The REporting of studies Conducted using Observational Routinely-collected health Data (RECORD) statement. PLoS Med. 2015;12(10):e1001885.",
    "Collins et al. 2015": "Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent reporting of a multivariable prediction model for individual prognosis or diagnosis (TRIPOD). Ann Intern Med. 2015;162(1):55-63.",
    "Moran 1950": "Moran PAP. Notes on continuous stochastic phenomena. Biometrika. 1950;37(1-2):17-23.",
    "Openshaw 1984": "Openshaw S. The modifiable areal unit problem. CATMOG 38. Norwich: Geo Books; 1984.",
    "Lundberg & Lee 2017": "Lundberg SM, Lee SI. A unified approach to interpreting model predictions. Adv Neural Inf Process Syst. 2017;30:4765-74.",
    "Alkire & Foster 2011": "Alkire S, Foster J. Counting and multidimensional poverty measurement. J Public Econ. 2011;95(7-8):476-87.",
    "Ghana Statistical Service 2021": "Ghana Statistical Service. Ghana 2021 Population and Housing Census: General Report. Accra: GSS; 2021.",
}

# NOTE: figure/table NUMBERS below reflect first-citation order in Results (renumbered
# 2026-07-01 after a user-caught defect: PNG filenames retain their original Stage-6 build
# order and do NOT match the final citation numbering -- e.g. "Figure 3" below uses the file
# figure6_mdr_gantt.png. This is intentional; do not "fix" filenames to match, just keep this
# dict as the single source of truth for the numbering.
FIGCAP = {
    1: ("figure1_incidence_trend.png", "Figure 1. National TB incidence, Ghana, 2000-2024, with Sen's slope trend and the 2020-2021 COVID-era window."),
    2: ("figure2_forecast_comparison.png", "Figure 2. Rolling-origin forecasts of national TB incidence by model (ARIMA, LSTM, XGBoost), Model A (full panel)."),
    3: ("figure6_mdr_gantt.png", "Figure 3. MDR-TB cascade series data-availability windows (locked, continuity-verified)."),
    4: ("figure3_tvi_choropleth.png", "Figure 4. District TB-Vulnerability Index (TVI-PCA), 261 districts, 2021 census cross-section."),
    5: ("figure4_lisa_clusters.png", "Figure 5. LISA clusters of district TVI-PCA (High-High, Low-Low, p<0.05)."),
    6: ("figure5_shap_importance.png", "Figure 6. SHAP mean |value| feature importance, XGBoost incidence forecast (3-year lag structure)."),
}

# NOTE: table NUMBERS reflect first-citation order in Results (renumbered 2026-07-01,
# same fix as FIGCAP above). CSV filenames retain their original Stage-6 names.
TABLES = {  # longest keys first for matching (single-digit keys still need this since "10" would
            # otherwise partial-match "1" -- not currently an issue at 5 tables, kept for safety)
    "1": ("table1a_national_tb_panel_summary.csv", "Table 1. Descriptive summary of the national TB panel (2000-2024).", None),
    "2": ("table2_mann_kendall_trends.csv", "Table 2. Mann-Kendall trend test results across all Track A series.", None),
    "3": ("table3_mdr_subanalysis.csv", "Table 3. MDR-TB cascade shorter-window sub-analysis (descriptive + one-step ARIMA forecast).", None),
    "4": ("table1b_district_vulnerability_summary.csv", "Table 4. Descriptive summary of district vulnerability-index components (n=261).", None),
    "5": ("table4_shap_importance.csv", "Table 5. SHAP feature importance, XGBoost incidence forecast.", None),
}


def load(f):
    return open(MAN + f, encoding="utf-8").read()


abstract = load("abstract.md")
im = load("draft_intro_methods.md")
res = load("draft_results.md")
dis = load("draft_discussion.md")
intro = im.split("## Methods")[0]
methods = "## Methods" + im.split("## Methods")[1]
title = re.search(r"^# (.+)$", im, re.M).group(1)
body_order = [("Introduction", intro), ("Methods", methods), ("Results", res), ("Discussion", dis)]
full_text = " ".join(t for _, t in body_order)

# ---- Vancouver re-index by first appearance ------------------------------------
pos = {k: full_text.find("(" + k) if ("(" + k) in full_text else min(
    [m.start() for m in re.finditer(re.escape(k), full_text)] or [10**9]) for k in REFS}
order = [k for k in sorted(REFS, key=lambda k: pos[k]) if pos[k] < 10**9]
num = {k: i + 1 for i, k in enumerate(order)}
paren_re = re.compile(r"\(([^()]*\d{4}[^()]*)\)")


def vancouver(txt):
    def repl(m):
        segs = [s.strip() for s in re.split(r";\s*", m.group(1))]
        if segs and all(s in num for s in segs):
            return "[" + ",".join(str(num[s]) for s in segs) + "]"
        return m.group(0)
    return paren_re.sub(repl, txt)


# ---- docx helpers ---------------------------------------------------------------
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)


def H(t, l=1):
    return doc.add_heading(t, level=l)


def rich(p, para):
    for seg in re.split(r"(\*\*.*?\*\*)", para):
        r = p.add_run(re.sub(r"\*\*", "", seg))
        if seg.startswith("**"):
            r.bold = True


def add_table(key):
    fn, cap, limit = TABLES[key]
    path = TAB + fn
    if not os.path.exists(path):
        return
    df = pd.read_csv(path)
    if limit:
        df = df.head(limit)
    df = relabel_cols(df)
    cp = doc.add_paragraph(cap)
    cp.runs[0].italic = True
    cp.runs[0].font.size = Pt(9)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light List Accent 1"
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(c)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, v in enumerate(row):
            sval = "" if (pd.isna(v) or str(v).lower() == "nan") else str(v)
            cells[j].text = sval
            if cells[j].paragraphs[0].runs:
                cells[j].paragraphs[0].runs[0].font.size = Pt(8)


# ---- title / front matter --------------------------------------------------------
doc.add_heading(title, 0)
doc.add_paragraph("Valentine Golden Ghanem").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Correspondence: valentineghanem@gmail.com").alignment = WD_ALIGN_PARAGRAPH.CENTER
rh = doc.add_paragraph()
rh.add_run("Running head: ").bold = True
rh.add_run("TB trends, treatment gaps and LSTM forecasting in Ghana")

H("Abstract", 1)
for para in abstract.split("\n\n"):
    para = para.strip()
    if not para or para.startswith("# Abstract"):
        continue
    rich(doc.add_paragraph(), para.replace("**Keywords:**", "Keywords:"))
ab = doc.add_paragraph()
ab.add_run("Abbreviations: ").bold = True
ab.add_run(
    "ART, antiretroviral therapy; ARIMA, autoregressive integrated moving average; AIC, Akaike information "
    "criterion; CI, confidence interval; DHIMS-2, District Health Information Management System 2; GDP, gross "
    "domestic product; GGHE-D, domestic general government health expenditure; GSS, Ghana Statistical Service; "
    "LISA, local indicators of spatial association; LSTM, long short-term memory; MAE, mean absolute error; MAPE, "
    "mean absolute percentage error; MDR-TB, multidrug-resistant tuberculosis; NTBLCP, National Tuberculosis and "
    "Leprosy Control Programme; PLHIV, people living with HIV; RMSE, root mean squared error; RR-TB, "
    "rifampicin-resistant tuberculosis; SHAP, SHapley Additive exPlanations; TVI, TB-Vulnerability Index; WHO-GHO, "
    "World Health Organization Global Health Observatory; XDR-TB, extensively drug-resistant tuberculosis.")

# ---- body (figures + tables embedded at first citation) --------------------------
fig_done = set()
tab_done = set()
for secname, txt in body_order:
    txt = vancouver(txt)
    H(secname, 1)
    for block in txt.split("\n\n"):
        block = block.strip()
        if not block or block.startswith("# "):
            continue
        if block.startswith("*") and not block.startswith("**"):
            continue  # skip italic metadata lines (e.g. "*Working draft...*"), not bold-led paragraphs
        para = block
        if block.startswith("## ") or block.startswith("### "):
            lines = block.split("\n", 1)
            sub = lines[0].lstrip("# ").strip()
            if sub.lower() not in (secname.lower(), "methods"):
                H(sub, 2)
            if len(lines) < 2 or not lines[1].strip():
                continue
            para = lines[1].strip()
        rich(doc.add_paragraph(), para)
        for n, (fn, cap) in FIGCAP.items():
            if f"Figure {n}" in para and n not in fig_done and os.path.exists(FIG + fn):
                doc.add_picture(FIG + fn, width=Inches(5.4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                c = doc.add_paragraph(cap)
                c.runs[0].italic = True
                c.runs[0].font.size = Pt(9)
                fig_done.add(n)
        for key in sorted(TABLES, key=len, reverse=True):
            if f"Table {key}" in para and key not in tab_done:
                add_table(key)
                tab_done.add(key)

# ---- declarations / back matter ----------------------------------------------------
H("Declarations", 1)


def decl(h, b):
    p = doc.add_paragraph()
    p.add_run(h + " ").bold = True
    p.add_run(b)


decl("Ethics approval and consent.", "This study analyses publicly available, de-identified, aggregate secondary data (WHO Global Health Observatory, the WHO Global Tuberculosis Report 2024, and Ghana Statistical Service 2021 Population and Housing Census district tables). No individual-level records, human participants, or identifiable information are involved. The work qualifies for exemption from full ethical review under Ghana Health Service Ethics Review Committee guidance on secondary analysis of anonymised public data.")
decl("Data availability.", "All source datasets are public (WHO Global Health Observatory, WHO Global Tuberculosis Report 2024, and Ghana Statistical Service census tables). Derived master datasets, locked hyperparameters, and all analysis code are available in the project repository.")
decl("Code availability.", "Python pipeline (build_master, build_vulnerability_index, lock_hyperparameters, mann_kendall_trends, forecast_rolling_origin, mdr_subanalysis, spatial_analysis, shap_decomposition, build_figures) released in the repository; reproducible end-to-end.")
decl("Funding.", "This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.")
decl("Competing interests.", "The author declares no competing interests.")
decl("Author contributions (CRediT).", "Valentine Golden Ghanem (ORCID: 0009-0002-8332-0220): conceptualization, methodology, software, formal analysis, data curation, visualization, writing — original draft, writing — review & editing.")
decl("Acknowledgements.", "The author thanks the WHO Global Health Observatory and the Ghana Statistical Service for open data access.")
decl("Reporting.", "The study is reported per STROBE and RECORD; the forecasting components follow TRIPOD-adapted reporting items. Completed checklists are provided as Supplementary Files.")
decl("Supplementary material.", "S1 STROBE checklist; S2 RECORD checklist; S3 forecast-tier metadata and locked hyperparameters; S4 supplementary figures and full LISA/Gi* district results.")

# ---- references -----------------------------------------------------------------
H("References", 1)
for k in order:
    doc.add_paragraph(f"{num[k]}. {REFS[k]}")

out = MAN + "TB_Trends_LSTM_Forecasting_Ghana_manuscript.docx"
doc.save(out)
print("Saved:", out)
print(f"Refs: {len(order)} | figures: {sorted(fig_done)} | tables: {sorted(tab_done)}")
uncited = [k for k in REFS if k not in num]
if uncited:
    print("WARNING uncited refs (dropped):", uncited)
